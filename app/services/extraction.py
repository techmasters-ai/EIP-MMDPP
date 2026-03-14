"""Document content extraction services.

Handles text, images, tables, OCR, and schematics from documents.
All processing runs locally (air-gapped deployment).
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ExtractedChunk:
    """A piece of extracted content ready for embedding."""

    chunk_text: str
    modality: str  # text | image | table | schematic | ocr
    page_number: Optional[int] = None
    bounding_box: Optional[dict] = None
    raw_image_bytes: Optional[bytes] = None  # for image/schematic chunks
    ocr_confidence: Optional[float] = None
    ocr_engine: Optional[str] = None
    requires_human_review: bool = False
    metadata: dict = field(default_factory=dict)


def extract_pdf(
    pdf_bytes: bytes,
    tesseract_threshold: float = 0.75,
    easyocr_threshold: float = 0.60,
) -> list[ExtractedChunk]:
    """Extract all content from a PDF file.

    Extraction strategy:
    - Text layers: pdfplumber (layout-aware)
    - Embedded images: pymupdf (fitz)
    - OCR on image regions: pytesseract → easyocr fallback
    - Confidence < easyocr_threshold: flag for human review
    """
    import pdfplumber
    import fitz  # pymupdf

    chunks: list[ExtractedChunk] = []

    # --- Text extraction via pdfplumber ---
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if text and text.strip():
                chunks.append(
                    ExtractedChunk(
                        chunk_text=text.strip(),
                        modality="text",
                        page_number=page_num,
                    )
                )

            # Extract tables as text
            for table in page.extract_tables():
                if table:
                    table_text = _table_to_text(table)
                    if table_text.strip():
                        chunks.append(
                            ExtractedChunk(
                                chunk_text=table_text,
                                modality="table",
                                page_number=page_num,
                            )
                        )

    # --- Image extraction via pymupdf ---
    fitz_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page_num in range(len(fitz_doc)):
        page = fitz_doc[page_num]
        image_list = page.get_images(full=True)

        for img_info in image_list:
            xref = img_info[0]
            base_image = fitz_doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]

            # Attempt OCR on extracted image
            ocr_chunks = _ocr_image(
                image_bytes,
                page_num=page_num + 1,
                tesseract_threshold=tesseract_threshold,
                easyocr_threshold=easyocr_threshold,
            )
            chunks.extend(ocr_chunks)

            # Also store the image itself for visual embedding
            chunks.append(
                ExtractedChunk(
                    chunk_text="",  # will be caption or empty
                    modality="image",
                    page_number=page_num + 1,
                    raw_image_bytes=image_bytes,
                    metadata={"ext": image_ext},
                )
            )

    fitz_doc.close()
    return chunks


def extract_docx(docx_bytes: bytes) -> list[ExtractedChunk]:
    """Extract text content from a DOCX file."""
    from docx import Document

    chunks: list[ExtractedChunk] = []
    doc = Document(io.BytesIO(docx_bytes))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        # Join into a single text chunk per logical section
        text = "\n".join(paragraphs)
        chunks.append(ExtractedChunk(chunk_text=text, modality="text"))

    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        table_text = _table_to_text(rows)
        if table_text.strip():
            chunks.append(ExtractedChunk(chunk_text=table_text, modality="table"))

    return chunks


def extract_txt(file_bytes: bytes) -> list[ExtractedChunk]:
    """Extract text content from a plain text file.

    Decodes as UTF-8 (falling back to latin-1), then splits into
    paragraph-based chunks suitable for embedding.
    """
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")

    text = text.strip()
    if not text:
        return []

    # Split on double newlines into paragraphs, then re-join small paragraphs
    # into chunks that respect the embedding token window.
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not paragraphs:
        return [ExtractedChunk(chunk_text=text, modality="text")]

    chunks: list[ExtractedChunk] = []
    current: list[str] = []
    current_len = 0
    max_chars = 2000  # ~512 tokens at ~4 chars/token

    for para in paragraphs:
        if current_len + len(para) > max_chars and current:
            chunks.append(
                ExtractedChunk(chunk_text="\n\n".join(current), modality="text")
            )
            current = []
            current_len = 0
        current.append(para)
        current_len += len(para)

    if current:
        chunks.append(
            ExtractedChunk(chunk_text="\n\n".join(current), modality="text")
        )

    return chunks


def extract_image(
    image_bytes: bytes,
    tesseract_threshold: float = 0.75,
    easyocr_threshold: float = 0.60,
) -> list[ExtractedChunk]:
    """Extract content from a standalone image file (PNG, JPG, TIFF)."""
    chunks = _ocr_image(
        image_bytes,
        tesseract_threshold=tesseract_threshold,
        easyocr_threshold=easyocr_threshold,
    )
    # Always include the image itself for visual embedding
    chunks.append(
        ExtractedChunk(
            chunk_text="",
            modality="image",
            raw_image_bytes=image_bytes,
        )
    )
    return chunks


def _ocr_image(
    image_bytes: bytes,
    page_num: Optional[int] = None,
    tesseract_threshold: float = 0.75,
    easyocr_threshold: float = 0.60,
) -> list[ExtractedChunk]:
    """OCR an image. Tries pytesseract first, falls back to easyocr.

    Low-confidence results are flagged for human review.
    """
    import pytesseract
    from PIL import Image

    chunks: list[ExtractedChunk] = []
    pil_image = Image.open(io.BytesIO(image_bytes))

    # --- pytesseract (primary) ---
    try:
        data = pytesseract.image_to_data(
            pil_image, output_type=pytesseract.Output.DICT, config="--psm 6"
        )
        confidences = [
            c for c in data["conf"] if isinstance(c, (int, float)) and c >= 0
        ]
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
        normalized_confidence = avg_confidence / 100.0  # tesseract uses 0-100

        text = " ".join(w for w in data["text"] if w.strip())

        if text.strip() and normalized_confidence >= tesseract_threshold:
            chunks.append(
                ExtractedChunk(
                    chunk_text=text.strip(),
                    modality="ocr",
                    page_number=page_num,
                    ocr_confidence=normalized_confidence,
                    ocr_engine="tesseract",
                )
            )
            return chunks

        # Confidence below tesseract threshold — try easyocr
        if text.strip() and normalized_confidence >= easyocr_threshold:
            # Use easyocr as a fallback
            easyocr_chunks = _easyocr_image(
                image_bytes, page_num, easyocr_threshold
            )
            if easyocr_chunks:
                return easyocr_chunks

        # Below both thresholds — flag for human review
        if text.strip():
            chunks.append(
                ExtractedChunk(
                    chunk_text=text.strip(),
                    modality="ocr",
                    page_number=page_num,
                    ocr_confidence=normalized_confidence,
                    ocr_engine="tesseract",
                    requires_human_review=True,
                )
            )

    except Exception as e:
        logger.warning("pytesseract failed: %s. Falling back to easyocr.", e)
        return _easyocr_image(image_bytes, page_num, easyocr_threshold)

    return chunks


def _easyocr_image(
    image_bytes: bytes,
    page_num: Optional[int] = None,
    confidence_threshold: float = 0.60,
) -> list[ExtractedChunk]:
    """Run easyocr on an image. Returns extracted chunks."""
    import easyocr
    import numpy as np
    from PIL import Image

    reader = easyocr.Reader(["en"], gpu=False, verbose=False)
    img_array = np.array(Image.open(io.BytesIO(image_bytes)))
    results = reader.readtext(img_array)

    if not results:
        return []

    texts = []
    confidences = []
    for _, text, conf in results:
        if text.strip():
            texts.append(text.strip())
            confidences.append(float(conf))

    if not texts:
        return []

    avg_confidence = float(sum(confidences) / len(confidences))
    combined_text = " ".join(texts)
    requires_review = avg_confidence < confidence_threshold

    return [
        ExtractedChunk(
            chunk_text=combined_text,
            modality="ocr",
            page_number=page_num,
            ocr_confidence=avg_confidence,
            ocr_engine="easyocr",
            requires_human_review=requires_review,
        )
    ]


def _table_to_text(table: list[list[str | None]]) -> str:
    """Convert a 2D table to a markdown-style text representation."""
    if not table:
        return ""
    rows = []
    for row in table:
        cells = [str(c or "").strip() for c in row]
        rows.append(" | ".join(cells))
    return "\n".join(rows)
